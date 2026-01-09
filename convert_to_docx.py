#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_rtl(paragraph):
    """Set paragraph to RTL direction"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set RTL direction
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def add_rtl_run(paragraph, text, font_size=None, bold=False, italic=False):
    """Add RTL text run to paragraph"""
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(font_size) if font_size else Pt(14)
    run.bold = bold
    run.italic = italic
    return run

def clean_markdown(text):
    """Remove markdown formatting but keep text"""
    # Remove markdown links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Remove markdown bold but keep text
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove markdown italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove code blocks
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

def parse_markdown_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set document page size
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        original_line = line
        line = line.strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Skip horizontal rules
        if line.startswith('---'):
            i += 1
            continue
        
        # H1 - Font size 16
        if line.startswith('# ') and not line.startswith('##'):
            text = clean_markdown(line[2:].strip())
            p = doc.add_paragraph()
            set_rtl(p)
            add_rtl_run(p, text, font_size=16, bold=True)
        
        # H2 - Font size 15
        elif line.startswith('## ') and not line.startswith('###'):
            text = clean_markdown(line[3:].strip())
            p = doc.add_paragraph()
            set_rtl(p)
            add_rtl_run(p, text, font_size=15, bold=True)
        
        # H3 - Font size 14, bold
        elif line.startswith('### '):
            text = clean_markdown(line[4:].strip())
            p = doc.add_paragraph()
            set_rtl(p)
            add_rtl_run(p, text, font_size=14, bold=True)
        
        # H4 - Font size 14, bold
        elif line.startswith('#### '):
            text = clean_markdown(line[5:].strip())
            p = doc.add_paragraph()
            set_rtl(p)
            add_rtl_run(p, text, font_size=14, bold=True)
        
        # Bold text (requirements) - lines starting with **
        elif line.startswith('**') and '**' in line[2:]:
            # Extract bold text
            match = re.match(r'\*\*(.*?)\*\*', line)
            if match:
                text = clean_markdown(match.group(1))
                remaining = clean_markdown(line[match.end():].strip())
                p = doc.add_paragraph()
                set_rtl(p)
                add_rtl_run(p, text, font_size=14, bold=True)
                if remaining:
                    add_rtl_run(p, remaining, font_size=14)
            else:
                text = clean_markdown(line)
                p = doc.add_paragraph()
                set_rtl(p)
                add_rtl_run(p, text, font_size=14, bold=True)
        
        # List items (including indented)
        elif line.startswith('- ') or (line.startswith('  - ') and not line.startswith('    -')):
            # Count indentation
            indent_level = 0
            if line.startswith('  - '):
                indent_level = 1
                text = clean_markdown(line[4:].strip())
            else:
                text = clean_markdown(line[2:].strip())
            
            p = doc.add_paragraph(style='List Bullet')
            set_rtl(p)
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            add_rtl_run(p, text, font_size=14)
        
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            text = clean_markdown(re.sub(r'^\d+\. ', '', line))
            p = doc.add_paragraph(style='List Number')
            set_rtl(p)
            add_rtl_run(p, text, font_size=14)
        
        # Regular text
        else:
            text = clean_markdown(line)
            # Handle placeholder tags
            if '<we-will' in text:
                text = text.replace('<we-will-add-use-case-diagram-here>', '[سيتم إضافة مخطط حالات الاستخدام هنا]')
                text = text.replace('<we-will-replace-this-with-class-diagram>', '[سيتم استبدال هذا بمخطط الفئات]')
            
            if text:
                p = doc.add_paragraph()
                set_rtl(p)
                add_rtl_run(p, text, font_size=14)
        
        i += 1
    
    doc.save(docx_file)
    print(f"تم إنشاء الملف: {docx_file}")

if __name__ == '__main__':
    md_file = 'مواصفات_متطلبات_البرمجيات.md'
    docx_file = 'مواصفات_متطلبات_البرمجيات.docx'
    parse_markdown_to_docx(md_file, docx_file)
